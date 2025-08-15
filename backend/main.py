from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Depends, status
from fastapi.responses import JSONResponse, Response
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from fastapi.security import OAuth2PasswordBearer, OAuth2PasswordRequestForm
from pydantic import BaseModel, EmailStr
from typing import List
import pdfplumber
import tempfile
import os
import re
from openai import OpenAI
from dotenv import load_dotenv
import motor.motor_asyncio
from datetime import datetime, timedelta
from passlib.context import CryptContext
import jwt
import base64
from io import BytesIO
from pdf2image import convert_from_path
import secrets
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import logging
import gridfs
from bson import ObjectId
# Load environment variables from .env file
load_dotenv()

main_app = FastAPI()
app = FastAPI()
app.mount("/backend", main_app)
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:5173",
        "https://final-resume-screening-app.vercel.app",
        "http://localhost:4173",
        "https://prohire.probusinsurance.com"
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)



# MongoDB setup
MONGODB_URI =os.getenv("MONGODB_URI")
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)
mongo_client = motor.motor_asyncio.AsyncIOMotorClient(MONGODB_URI)
db = mongo_client["resume_screening"]
mis_collection = db["mis"]
recruiters_collection = db["recruiters"]
reset_tokens_collection = db["reset_tokens"]  # New collection for reset tokens
fs = motor.motor_asyncio.AsyncIOMotorGridFSBucket(db)
# JWT setup
SECRET_KEY ="supersecretkey"
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = 60 * 24 * 7  # 1 week
RESET_TOKEN_EXPIRE_MINUTES = 30  # 30 minutes for reset tokens

# Email configuration
SMTP_SERVER = os.getenv("SMTP_SERVER")
SMTP_PORT = os.getenv("SMTP_PORT")
EMAIL_USERNAME=os.getenv("EMAIL_USERNAME")
EMAIL_PASSWORD=os.getenv("EMAIL_PASSWORD")
FROM_EMAIL=os.getenv("FROM_EMAIL")
FROM_NAME=os.getenv("FROM_NAME")

pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")
oauth2_scheme = OAuth2PasswordBearer(tokenUrl="/login")

# Make sure to set OPENAI_API_KEY in your .env file
OPENAI_API_KEY=os.getenv("OPENAI_API_KEY")
client = OpenAI(api_key=OPENAI_API_KEY)

# Pydantic models for request/response
class ForgotPasswordRequest(BaseModel):
    email: EmailStr

class ResetPasswordRequest(BaseModel):
    token: str
    new_password: str

class RecruiterRegistration(BaseModel):
    username: str
    password: str
    email: EmailStr

# --- Email Helper Functions ---
async def send_email(to_email: str, subject: str, body: str, is_html: bool = False):
    """Send email using SMTP"""
    try:
        if not EMAIL_USERNAME or not EMAIL_PASSWORD:
            raise HTTPException(
                status_code=500, 
                detail="Email configuration not set up properly"
            )
        
        msg = MIMEMultipart()
        msg["From"] = f"{FROM_NAME} <{FROM_EMAIL}>"
        msg['To'] = to_email
        msg['Subject'] = subject
        
        msg.attach(MIMEText(body, 'html' if is_html else 'plain'))
        
        server = smtplib.SMTP(SMTP_SERVER, SMTP_PORT)
        server.starttls()
        server.login(EMAIL_USERNAME, EMAIL_PASSWORD)
        text = msg.as_string()
        server.sendmail(FROM_EMAIL, to_email, text)
        server.quit()
        
        return True
    except Exception as e:
        print(f"Failed to send email: {e}")
        return False

def generate_reset_token():
    """Generate a secure random token for password reset"""
    return secrets.token_urlsafe(32)

def create_reset_token(email: str):
    """Create a JWT token for password reset"""
    expire = datetime.utcnow() + timedelta(minutes=RESET_TOKEN_EXPIRE_MINUTES)
    to_encode = {"email": email, "exp": expire, "type": "reset"}
    return jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)

def verify_reset_token(token: str):
    """Verify and decode reset token"""
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        if payload.get("type") != "reset":
            return None
        return payload.get("email")
    except jwt.PyJWTError:
        return None

# --- Recruiter Auth Helpers ---
def verify_password(plain, hashed):
    return pwd_context.verify(plain, hashed)

def get_password_hash(password):
    return pwd_context.hash(password)

async def get_recruiter(username: str):
    return await recruiters_collection.find_one({"username": username})

async def get_recruiter_by_email(email: str):
    return await recruiters_collection.find_one({"email": email})

def create_access_token(data: dict, expires_delta: timedelta = None):
    to_encode = data.copy()
    expire = datetime.utcnow() + (expires_delta or timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES))
    to_encode.update({"exp": expire})
    return jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)

async def get_current_recruiter(token: str = Depends(oauth2_scheme)):
    credentials_exception = HTTPException(
        status_code=status.HTTP_401_UNAUTHORIZED,
        detail="Could not validate credentials",
        headers={"WWW-Authenticate": "Bearer"},
    )
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        username: str = payload.get("sub")
        if username is None:
            raise credentials_exception
    except jwt.PyJWTError:
        raise credentials_exception
    recruiter = await get_recruiter(username)
    if recruiter is None:
        raise credentials_exception
    return recruiter

# --- Auth Endpoints ---
@main_app.post("/register")
async def register(recruiter_data: RecruiterRegistration):
    username = recruiter_data.username.strip()
    email = recruiter_data.email.strip().lower()
    
    # Check if username already exists
    existing_username = await recruiters_collection.find_one({"username": username})
    if existing_username:
        raise HTTPException(status_code=400, detail="Username already registered")
    
    # Check if email already exists
    existing_email = await recruiters_collection.find_one({"email": email})
    if existing_email:
        raise HTTPException(status_code=400, detail="Email already registered")
    
    hashed = get_password_hash(recruiter_data.password)
    
    await recruiters_collection.insert_one({
        "username": username,
        "email": email,
        "hashed_password": hashed,
        "created_at": datetime.utcnow()
    })
    
    return {"msg": "Recruiter registered successfully"}

@main_app.post("/register-form")
async def register_form(form: OAuth2PasswordRequestForm = Depends()):
    """Legacy registration endpoint for form-based registration"""
    username = form.username.strip()
    existing = await recruiters_collection.find_one({"username": username})
    if existing:
        raise HTTPException(status_code=400, detail="Username already registered")
    hashed = get_password_hash(form.password)
    await recruiters_collection.insert_one({
        "username": form.username, 
        "hashed_password": hashed,
        "created_at": datetime.utcnow()
    })
    return {"msg": "Recruiter registered"}

@main_app.post("/login")
async def login(form: OAuth2PasswordRequestForm = Depends()):
    username = form.username.strip()
    recruiter = await recruiters_collection.find_one({"username": username})
    if not recruiter or not verify_password(form.password, recruiter["hashed_password"]):
        raise HTTPException(status_code=400, detail="Incorrect username or password")
    access_token = create_access_token(data={"sub": recruiter["username"]})
    return {
        "access_token": access_token, 
        "token_type": "bearer", 
        "recruiter_name": recruiter["username"]
    }

@main_app.post("/forgot-password")
async def forgot_password(request: ForgotPasswordRequest):
    """Send password reset email"""
    email = request.email.strip().lower()
    
    # Check if email exists
    recruiter = await get_recruiter_by_email(email)
    if not recruiter:
        # Don't reveal if email exists or not for security
        return {"msg": "If the email exists, you will receive a password reset link"}
    
    # Generate reset token
    reset_token = create_reset_token(email)
    
    # Store token in database with expiration
    await reset_tokens_collection.insert_one({
        "email": email,
        "token": reset_token,
        "created_at": datetime.utcnow(),
        "expires_at": datetime.utcnow() + timedelta(minutes=RESET_TOKEN_EXPIRE_MINUTES),
        "used": False
    })
    
    # Create reset link (adjust the frontend URL as needed)
    FRONTEND_BASE_URL =os.getenv("FRONTEND_BASE_URL")
    reset_link = f"{FRONTEND_BASE_URL}/reset-password?token={reset_token}"

    
    # Email content
    subject = "Password Reset Request - Prohire"
    body = f"""
    <html>
        <body>
            <h2>Password Reset Request</h2>
            <p>Hello {recruiter['username']},</p>
            <p>You have requested to reset your password for Prohire</p>
            <p>Click the link below to reset your password:</p>
            <p><a href="{reset_link}" style="background-color: #4CAF50; color: white; padding: 10px 20px; text-decoration: none; border-radius: 4px;">Reset Password</a></p>
            <p>This link will expire in {RESET_TOKEN_EXPIRE_MINUTES} minutes.</p>
            <p>If you didn't request this reset, please ignore this email.</p>
            <br>
            <p>Best regards,<br>ProHire Team</p>
        </body>
    </html>
    """
    
    # Send email
    email_sent = await send_email(email, subject, body, is_html=True)
    
    if not email_sent:
        raise HTTPException(status_code=500, detail="Failed to send reset email")
    
    return {"msg": "If the email exists, you will receive a password reset link"}

@main_app.post("/reset-password")
async def reset_password(request: ResetPasswordRequest):
    """Reset password using token"""
    
    # Verify token
    email = verify_reset_token(request.token)
    if not email:
        raise HTTPException(status_code=400, detail="Invalid or expired reset token")
    
    # Check if token exists in database and is not used
    token_doc = await reset_tokens_collection.find_one({
        "token": request.token,
        "used": False,
        "expires_at": {"$gt": datetime.utcnow()}
    })
    
    if not token_doc:
        raise HTTPException(status_code=400, detail="Invalid or expired reset token")
    
    # Get recruiter
    recruiter = await get_recruiter_by_email(email)
    if not recruiter:
        raise HTTPException(status_code=404, detail="User not found")
    
    # Update password
    hashed_password = get_password_hash(request.new_password)
    
    await recruiters_collection.update_one(
        {"email": email},
        {
            "$set": {
                "hashed_password": hashed_password,
                "password_updated_at": datetime.utcnow()
            }
        }
    )
    
    # Mark token as used
    await reset_tokens_collection.update_one(
        {"token": request.token},
        {"$set": {"used": True, "used_at": datetime.utcnow()}}
    )
    
    return {"msg": "Password reset successfully"}

@main_app.get("/verify-reset-token/{token}")
async def verify_reset_token_endpoint(token: str):
    """Verify if reset token is valid"""
    
    email = verify_reset_token(token)
    if not email:
        raise HTTPException(status_code=400, detail="Invalid or expired reset token")
    
    # Check if token exists in database and is not used
    token_doc = await reset_tokens_collection.find_one({
        "token": token,
        "used": False,
        "expires_at": {"$gt": datetime.utcnow()}
    })
    
    if not token_doc:
        raise HTTPException(status_code=400, detail="Invalid or expired reset token")
    
    return {"valid": True, "email": email}

# Cleanup expired tokens (call this periodically)
@main_app.delete("/cleanup-expired-tokens")
async def cleanup_expired_tokens():
    """Remove expired reset tokens from database"""
    result = await reset_tokens_collection.delete_many({
        "expires_at": {"$lt": datetime.utcnow()}
    })
    return {"deleted_count": result.deleted_count}

def extract_text_from_pdf(filepath):
    with pdfplumber.open(filepath) as pdf:
        extracted_text = '\n'.join(
            page.extract_text() for page in pdf.pages if page.extract_text()
        )
        if extracted_text.strip():
            return extracted_text.strip()

    # If no text was found using pdfplumber, fallback to OCR using OpenAI
    try:
        images = convert_from_path(filepath)
        full_ocr_text = []

        for img in images:
            buffered = BytesIO()
            img.save(buffered, format="PNG")
            img_base64 = base64.b64encode(buffered.getvalue()).decode()

            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/png;base64,{img_base64}"
                                }
                            },
                            {
                                "type": "text",
                                "text": "Please extract all readable text from this image of a resume."
                            }
                        ]
                    }
                ],
                temperature=0.3,
                max_tokens=1500
            )

            ocr_text = response.choices[0].message.content.strip()
            if ocr_text:
                full_ocr_text.append(ocr_text)

        return "\n".join(full_ocr_text) if full_ocr_text else "❌ No text found in image using OCR."

    except Exception as e:
        return f"❌ Error during OCR fallback: {e}"




def extract_text_from_docx(filepath: str) -> str:
    """
    Enhanced DOCX extractor for Naukri resumes and other structured documents.
    Tries structured paragraph and table parsing first, then falls back to docx2txt.
    """
    try:
        from docx import Document
        import docx2txt

        full_text = []

        # Primary extraction using python-docx
        doc = Document(filepath)

        # Extract paragraphs
        for para in doc.paragraphs:
            line = para.text.strip()
            if line and line not in full_text:
                full_text.append(line)

        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text and cell_text not in row_data:
                        row_data.append(cell_text)
                if row_data:
                    full_text.append(" | ".join(row_data))

        # Fallback to docx2txt if primary fails
        if not full_text or len("".join(full_text).strip()) < 30:
            fallback_text = docx2txt.process(filepath)
            if fallback_text and fallback_text.strip():
                full_text = [line.strip() for line in fallback_text.splitlines() if line.strip()]

        # Final cleanup: remove exact duplicates while preserving order
        final_text = "\n".join(dict.fromkeys(full_text))
        return final_text.strip() if final_text else "❌ Could not extract usable text from this resume."

    except Exception as e:
        return f"❌ Error extracting text from DOCX: {e}"


def extract_text_from_image(filepath: str) -> str:
    """
    Extract text from image files using OpenAI's Vision API (GPT-4 Vision).
    Supports common image formats like PNG, JPG, JPEG, etc.
    """
    try:
        print(f"Starting OCR processing for file: {filepath}")  # Debug log
        
        # Check if file exists
        if not os.path.exists(filepath):
            return f"❌ Error: File {filepath} not found"
        
        # Read the image file
        with open(filepath, "rb") as image_file:
            image_data = image_file.read()
            print(f"Image file size: {len(image_data)} bytes")  # Debug log
            
            # Encode to base64
            base64_image = base64.b64encode(image_data).decode('utf-8')
            print(f"Base64 encoding completed, length: {len(base64_image)}")  # Debug log
            
            # Use OpenAI's Vision API to extract text
            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": "Please extract all the text from this image. Return only the extracted text without any additional formatting or explanations."
                            },
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/jpeg;base64,{base64_image}"
                                }
                            }
                        ]
                    }
                ],
                max_tokens=1000,
                temperature=0.1
            )
            
            extracted_text = response.choices[0].message.content
            print(f"OCR completed, extracted text length: {len(extracted_text) if extracted_text else 0}")  # Debug log
            
            if extracted_text and extracted_text.strip():
                return extracted_text.strip()
            else:
                return "❌ Could not extract text from this image. Please ensure the image contains clear, readable text."
                
    except Exception as e:
        print(f"Error in OCR processing: {str(e)}")  # Debug log
        return f"❌ Error extracting text from image: {e}"




def analyze_resume(jd, resume_text, hiring_choice, level_choice):
    prompt = ""
    if hiring_choice == "1":
        if level_choice == "1":
            prompt = f"""
You are a professional HR assistant AI screening resumes for a **Sales Fresher** role.

--- Job Description ---
{jd}

--- Candidate Resume ---
{resume_text}

--- Screening Criteria ---
1. Location: 
   - Candidate must be either from the job location city (e.g., Kolkata) or nearby cities (e.g., Durgapur) within feasible travel distance.
   - If candidate is not in the exact city but lives in a nearby town and the job allows remote or field sales operations, they should be considered.
   - Candidate should be able to travel to the main office once a month for reporting.
2. Age: As per job description.
3. Education: 12th pass & above.
4. Gender: As per job description.

Note: Everything should match the Job Description.

--- Response Format ---
Match %: XX%
Pros:
- ...
Cons:
- ...
Decision: ✅ Shortlist or ❌ Reject
Reason (if Rejected): ...
"""
        elif level_choice == "2":
            prompt = f"""
You are a professional HR assistant AI screening resumes for a **Sales Experienced** role.

--- Job Description ---
{jd}

--- Candidate Resume ---
{resume_text}

--- Screening Criteria ---
1. Location: 
   - Candidate must be either from the job location city (e.g., Kolkata) or nearby cities (e.g., Durgapur) within feasible travel distance.
   - If candidate is not in the exact city but lives in a nearby town and the job allows remote or field sales operations, they should be considered.
   - Candidate should be able to travel to the main office once a month for reporting.
2. Age: As per job description ("up to" logic preferred).
3. Total Experience: Add all types of sales (health + motor, etc.).
4. Relevant Experience: Must match industry (strict).
5. Education: 12th pass & above accepted.
6. Gender: As per job description.
7. Skills: Skills should align with relevant experience.
8. Stability: Ignore if 1 job <1 year; Reject if 2+ jobs each <1 year.

Note: Everything should match the Job Description.

--- Response Format ---
Match %: XX%
Pros:
- ...
Cons:
- ...
Decision: ✅ Shortlist or ❌ Reject
Reason (if Rejected): ...
"""
    elif hiring_choice == "2":
        if level_choice == "1":
            prompt = f"""
You are a professional HR assistant AI screening resumes for an **IT Fresher** role.

--- Job Description ---
{jd}

--- Candidate Resume ---
{resume_text}

--- Screening Criteria ---
1. Location: Must be local.
2. Age: Ignore or as per JD.
3. Experience: Internship is a bonus; no experience is fine.
4. Projects: Highlighted as experience if relevant.
5. Education: B.E, M.E, BTech, MTech, or equivalent in IT.
6. Gender: As per job description.
7. Skills: Must align with the job field (e.g., Full Stack).
Note: For example, if hiring for a Full Stack Engineer role, even if one or two skills mentioned in the Job Description are missing, the candidate can still be considered if they have successfully built Full Stack projects. Additional skills or tools mentioned in the JD are good-to-have, but not mandatory.
8. Stability: Not applicable.

Note: Everything should match the Job Description.

--- Response Format ---
Match %: XX%
Pros:
- ...
Cons:
- ...
Decision: ✅ Shortlist or ❌ Reject
Reason (if Rejected): ...
"""
        elif level_choice == "2":
            prompt = f"""
You are a professional HR assistant AI screening resumes for an **IT Experienced** role.

--- Job Description ---
{jd}

--- Candidate Resume ---
{resume_text}

--- Screening Criteria ---
1. Location: Must be local.
2. Age: As per job description (prefer "up to").
3. Total Experience: Overall IT field experience.
4. Relevant Experience: Must align with JD field.
5. Education: IT-related degrees only (B.E, M.Tech, etc.).
6. Gender: As per job description.
7. Skills: Languages and frameworks should match JD.
8. Stability: Ignore if 1 company <1 year; Reject if 2+ companies each <1 year.

Note: Everything should match the Job Description.

--- Response Format ---
Match %: XX%
Pros:
- ...
Cons:
- ...
Decision: ✅ Shortlist or ❌ Reject
Reason (if Rejected): ...
"""
    elif hiring_choice == "3":
        if level_choice == "1":
            prompt = f"""
You are a professional HR assistant AI screening resumes for a **Non-Sales Fresher** role.

--- Job Description ---
{jd}

--- Candidate Resume ---
{resume_text}

--- Screening Criteria ---
1. Location: Should be local and match JD.
2. Age: As per JD.
3. Total / Relevant Experience: Internship optional, but candidate should have certifications.
4. Education: Must be relevant to the JD.
5. Gender: As per JD.
6. Skills: Must align with the JD.
7. Stability: Not applicable for freshers.

Note: Everything should match the Job Description.

--- Response Format ---
Match %: XX%
Pros:
- ...
Cons:
- ...
Decision: ✅ Shortlist or ❌ Reject
Reason (if Rejected): ...
"""
        elif level_choice == "2":
            prompt = f"""
You are a professional HR assistant AI screening resumes for a **Non-Sales Experienced** role.

--- Job Description ---
{jd}

--- Candidate Resume ---
{resume_text}

--- Screening Criteria ---
1. Location: Must strictly match the JD.
2. Age: As per JD.
3. Total Experience: Overall professional experience.
4. Relevant Experience: Must align with role in JD.
5. Education: Must match the JD.
6. Gender: As per JD.
7. Skills: Should align with JD and match relevant experience (skills = relevant experience).
8. Stability:
   - If 2+ companies and each job ≤1 year → Reject.
   - If 1 company and ≤1 year → Ignore stability.

Note: Everything should match the Job Description.

--- Response Format ---
Match %: XX%
Pros:
- ...
Cons:
- ...
Decision: ✅ Shortlist or ❌ Reject
Reason (if Rejected): ...
"""     
    elif hiring_choice == "4":
        if level_choice == "1":
            prompt = f"""
You are a professional HR assistant AI screening resumes for a **Sales Support Fresher** role.

--- Job Description ---
{jd}

--- Candidate Resume ---
{resume_text}

--- Screening Criteria ---
1. Location: Must be strictly local.
2. Age: As per job description.
3. Education: 12th pass & above.
4. Gender: As per job description.

Note: Everything should match the Job Description.

--- Response Format ---
Match %: XX%
Pros:
- ...
Cons:
- ...
Decision: ✅ Shortlist or ❌ Reject
Reason (if Rejected): ...
"""
        elif level_choice == "2":
            prompt = f"""
You are a professional HR assistant AI screening resumes for a **Sales Support Experienced** role.

--- Job Description ---
{jd}

--- Candidate Resume ---
{resume_text}

--- Screening Criteria ---
1. Location: Must be strictly local.
2. Age: As per job description ("up to" logic preferred).
3. Total Experience: Add all types of sales support
4. Relevant Experience: Must match industry (strict).
5. Education: 12th pass & above accepted.
6. Gender: As per job description.
7. Skills: Skills should align with relevant experience.
8. Stability: Ignore if 1 job <1 year; Reject if 2+ jobs each <1 year.

Note: Everything should match the Job Description.

--- Response Format ---
Match %: XX%
Pros:
- ...
Cons:
- ...
Decision: ✅ Shortlist or ❌ Reject
Reason (if Rejected): ...
"""

    if not prompt:
        return "❌ Error: Invalid hiring or level choice provided."

    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.3,
        max_tokens=800
    )
    content = response.choices[0].message.content
    if content is not None:
        result_text = content.strip()
    else:
        result_text = "Match %: 0\nDecision: ❌ Reject\nReason (if Rejected): No response from model."
    usage = getattr(response, 'usage', None)

    match_percent = 0
    match_line = re.search(r"Match\s*%:\s*(\d+)", result_text)
    if match_line:
        match_percent = int(match_line.group(1))

    if match_percent < 72:
        result_text = re.sub(r"Decision:.*", "Decision: ❌ Reject", result_text)
        if "Reason (if Rejected):" in result_text:
            result_text = re.sub(r"Reason \(if Rejected\):.*", "Reason (if Rejected): Match % below 72% threshold.", result_text)
        else:
            result_text += "\nReason (if Rejected): Match % below 72% threshold."

    return {
        "result_text": result_text,
        "match_percent": match_percent,
        "usage": {
            "prompt_tokens": getattr(usage, 'prompt_tokens', None),
            "completion_tokens": getattr(usage, 'completion_tokens', None),
            "total_tokens": getattr(usage, 'total_tokens', None)
        } if usage else None
    }

def extract_candidate_name(resume_text, filename):
    # This function is now unused, but kept for reference
    return ""

def get_hiring_type_label(hiring_type):
    return {"1": "Sales", "2": "IT", "3": "Non-Sales","4":"Sales Support"}.get(hiring_type, hiring_type)

def get_level_label(level):
    return {"1": "Fresher", "2": "Experienced"}.get(level, level)

def format_date_with_day(dt):
    """Format date as '25th August 2025, Monday'"""
    day = dt.day
    if 10 <= day % 100 <= 20:
        suffix = 'th'
    else:
        suffix = {1: 'st', 2: 'nd', 3: 'rd'}.get(day % 10, 'th')
    
    formatted_date = dt.strftime(f"%d{suffix} %B %Y, %A")
    return formatted_date

@main_app.post("/analyze-resumes/")
async def analyze_resumes(
    job_description: str = Form(...),
    hiring_type: str = Form(...),
    level: str = Form(...),
    files: List[UploadFile] = File(...),
    recruiter=Depends(get_current_recruiter)
):
    results = []
    shortlisted = 0
    rejected = 0
    history = []
    current_date = datetime.utcnow()
    hiring_type_label = get_hiring_type_label(hiring_type)
    level_label = get_level_label(level)
    
    for file in files:
        filename = file.filename or "Unknown"
        suffix = os.path.splitext(filename)[1].lower()
        print(f"Processing file: {filename} with suffix: {suffix}")  # Debug log
        
        # Define supported image extensions
        supported_images = [".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tiff", ".webp"]
        
        # Read file content once
        file_content = await file.read()
        
        # Store file in GridFS regardless of type
        file_id = None
        try:
            file_id = await fs.upload_from_stream(
                filename,
                file_content,
                metadata={
                    "content_type": file.content_type or "application/octet-stream",
                    "upload_date": current_date,
                    "recruiter_name": recruiter["username"],
                     "file_size": len(file_content)
                }
            )

            print(f"File stored in GridFS with ID: {file_id}")
        except Exception as e:
            print(f"Failed to store file in GridFS: {e}")
        
        # Create temporary file for processing
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(file_content)
            tmp_path = tmp.name
            
        # Process different file types
        if suffix == ".pdf":
            resume_text = extract_text_from_pdf(tmp_path)
        elif suffix == ".docx":
            resume_text = extract_text_from_docx(tmp_path)
        elif suffix in supported_images:
            print(f"Processing image file: {filename} with suffix: {suffix}")  # Debug log
            resume_text = extract_text_from_image(tmp_path)
        else:
            os.unlink(tmp_path)
            error_msg = f"Unsupported file type: {suffix}. Only PDF, DOCX, and image files (JPG, JPEG, PNG, GIF, BMP, TIFF, WEBP) are allowed."
            print(f"File rejected: {filename} with suffix: {suffix}")  # Debug log
            results.append({
                "filename": filename,
                "error": error_msg
            })
            history.append({
                "resume_name": filename,
                "hiring_type": hiring_type_label,
                "level": level_label,
                "match_percent": None,
                "decision": "Error",
                "details": error_msg,
                "upload_date": format_date_with_day(current_date),
                "file_id": str(file_id) if file_id else None
            })
            continue
        
        # Analyze resume
        analysis = analyze_resume(job_description, resume_text, hiring_type, level)
        if isinstance(analysis, dict):
            analysis["filename"] = filename
            # Decision extraction
            decision = analysis.get("decision")
            if not decision and analysis.get("result_text"):
                match = re.search(r"Decision:\s*(✅ Shortlist|❌ Reject)", analysis["result_text"])
                if match:
                    decision = match.group(1)
            if decision and "Shortlist" in decision:
                shortlisted += 1
            elif decision and "Reject" in decision:
                rejected += 1
            decision_label = ("Shortlisted" if decision and "Shortlist" in decision else
                              "Rejected" if decision and "Reject" in decision else "-")
            analysis["decision"] = decision_label
            results.append(analysis)
            history.append({
                "resume_name": filename,
                "hiring_type": hiring_type_label,
                "level": level_label,
                "match_percent": analysis.get("match_percent"),
                "decision": decision_label,
                "details": analysis.get("result_text") or analysis.get("error", ""),
                "upload_date": format_date_with_day(current_date),
                "file_id": str(file_id) if file_id else None
            })
        else:
            results.append({"filename": filename, "error": analysis})
            history.append({
                "resume_name": filename,
                "hiring_type": hiring_type_label,
                "level": level_label,
                "match_percent": None,
                "decision": "Error",
                "details": analysis,
                "upload_date": format_date_with_day(current_date),
                "file_id": str(file_id) if file_id else None
            })
        
        # Clean up temporary file
        os.unlink(tmp_path)
    
    # Save MIS record with history
    await mis_collection.insert_one({
        "recruiter_name": recruiter["username"],
        "total_resumes": len(files),
        "shortlisted": shortlisted,
        "rejected": rejected,
        "timestamp": current_date,
        "history": history
    })
    return JSONResponse(content={"results": results})

@main_app.get("/download-resume/{file_id}")
async def download_resume(file_id: str, recruiter=Depends(get_current_recruiter)):
    try:
        grid_out = await fs.open_download_stream(ObjectId(file_id))
        file_content = await grid_out.read()
        
        return Response(
            content=file_content,
            media_type=grid_out.metadata.get("content_type", "application/octet-stream"),
            headers={
                "Content-Disposition": f"attachment; filename={grid_out.filename}"
            }
        )
    except Exception as e:
        raise HTTPException(status_code=404, detail="File not found")

@main_app.get("/view-resume/{file_id}")
async def view_resume(file_id: str, recruiter=Depends(get_current_recruiter)):
    try:
        grid_out = await fs.open_download_stream(ObjectId(file_id))
        file_content = await grid_out.read()
        
        return {
            "filename": grid_out.filename,
            "content_type": grid_out.metadata.get("content_type", "application/octet-stream"),
            "size": len(file_content),
            "content": base64.b64encode(file_content).decode('utf-8')
        }
    except Exception as e:
        raise HTTPException(status_code=404, detail="File not found")

@main_app.get("/mis-summary")
async def mis_summary():
    pipeline = [
        {
            "$group": {
                "_id": "$recruiter_name",
                "uploads": {"$sum": 1},
                "total_resumes": {"$sum": "$total_resumes"},
                "shortlisted": {"$sum": "$shortlisted"},
                "rejected": {"$sum": "$rejected"},
                "history": {"$push": "$history"}
            }
        },
        {"$sort": {"_id": 1}}
    ]
    
    summary = []
    async for row in mis_collection.aggregate(pipeline):
        # Flatten the list of lists in history
        flat_history = [item for sublist in row["history"] for item in sublist]
        
        # Calculate daily counts for this recruiter
        daily_counts = {}
        for item in flat_history:
            upload_date = item.get("upload_date", "")
            if upload_date:
                # Extract just the date part (without day name) for counting
                date_part = upload_date.split(',')[0] if ',' in upload_date else upload_date
                daily_counts[date_part] = daily_counts.get(date_part, 0) + 1
        
        # Add daily count to each history item
        for item in flat_history:
            upload_date = item.get("upload_date", "")
            if upload_date:
                date_part = upload_date.split(',')[0] if ',' in upload_date else upload_date
                item["counts_per_day"] = daily_counts.get(date_part, 0)
            else:
                item["counts_per_day"] = 0
        
        summary.append({
            "recruiter_name": row["_id"],
            "uploads": row["uploads"],
            "resumes": row["total_resumes"],
            "shortlisted": row["shortlisted"],
            "rejected": row["rejected"],
            "history": flat_history
        })
    return {"summary": summary}

@main_app.get("/daily-reports")
async def daily_reports():
    # Get today's date in UTC
    today = datetime.utcnow().replace(hour=0, minute=0, second=0, microsecond=0)
    tomorrow = today + timedelta(days=1)
    
    # Aggregate data for today only
    pipeline = [
        {
            "$match": {
                "timestamp": {
                    "$gte": today,
                    "$lt": tomorrow
                }
            }
        },
        {
            "$group": {
                "_id": "$recruiter_name",
                "total_resumes": {"$sum": "$total_resumes"},
                "shortlisted": {"$sum": "$shortlisted"},
                "rejected": {"$sum": "$rejected"}
            }
        },
        {"$sort": {"_id": 1}}
    ]
    
    daily_data = []
    async for row in mis_collection.aggregate(pipeline):
        daily_data.append({
            "recruiter_name": row["_id"],
            "total_resumes": row["total_resumes"],
            "shortlisted": row["shortlisted"],
            "rejected": row["rejected"]
        })
    
    # Format today's date
    today_formatted = format_date_with_day(today)
    
    return {
        "date": today_formatted,
        "reports": daily_data
    }

@main_app.get("/health")
async def health():
    return {"status": "ok"}

@main_app.get("/")
async def root():
    return {"message": "Backend is live!"}

@main_app.get("/")
async def backend_root():
    return {"message": "Backend API is live!"}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)